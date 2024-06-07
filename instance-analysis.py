import os
import requests
from requests.auth import HTTPBasicAuth
import logging
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s: %(message)s')

# Configuration for both Jira instances
source_config = {
    'email': 'rodolfobortolin@gmail.com',
    'token': '',
    'base_url': 'https://source.atlassian.net'
}

target_config = {
    'email': 'rodolfobortolin@gmail.com',
    'token': '',
    'base_url': 'https://target.atlassian.net'
}

# Function to get data from Jira instance
def get_jira_data(config, endpoint):
    url = f"{config['base_url']}{endpoint}"
    auth = HTTPBasicAuth(config['email'], config['token'])
    headers = {"Accept": "application/json"}
    
    response = requests.get(url, auth=auth, headers=headers)
    response.raise_for_status()
    return response.json()

# Function to search filters with pagination
def search_filters(config):
    filters = []
    start_at = 0
    max_results = 50
    while True:
        url = f"{config['base_url']}/rest/api/2/filter/search?expand=description,owner,jql,sharePermissions,editPermissions&startAt={start_at}&maxResults={max_results}"
        response = requests.get(url, headers={"Accept": "application/json"}, auth=HTTPBasicAuth(config['email'], config['token']))
        response.raise_for_status()
        data = response.json()
        filters.extend(data.get('values', []))
        if data.get('isLast', True):
            break
        start_at += max_results
    return filters

# Function to search dashboards with pagination
def search_dashboards(config):
    dashboards = []
    start_at = 0
    max_results = 50
    while True:
        url = f"{config['base_url']}/rest/api/2/dashboard?startAt={start_at}&maxResults={max_results}"
        response = requests.get(url, headers={"Accept": "application/json"}, auth=HTTPBasicAuth(config['email'], config['token']))
        response.raise_for_status()
        data = response.json()
        dashboards.extend(data.get('dashboards', []))
        if data.get('isLast', True):
            break
        start_at += max_results
    return [dashboard for dashboard in dashboards if dashboard['name'] != 'Default dashboard']

# Define endpoints for required data
endpoints = {
    'projects': '/rest/api/3/project',
    'priorities': '/rest/api/3/priority',
    'resolutions': '/rest/api/3/resolution',
    'roles': '/rest/api/3/role',
    'issuetypes': '/rest/api/3/issuetype',
    'customfields': '/rest/api/3/field',
    'statuses': '/rest/api/3/status',
    'filters': '/rest/api/2/filter/search',
    'dashboards': '/rest/api/2/dashboard/search'
}

# Fetch data from both instances
data_source = {key: get_jira_data(source_config, endpoint) for key, endpoint in endpoints.items()}
data_target = {key: get_jira_data(target_config, endpoint) for key, endpoint in endpoints.items()}

# Fetch filters and dashboards with pagination
data_source['filters'] = search_filters(source_config)
data_target['filters'] = search_filters(target_config)
data_source['dashboards'] = search_dashboards(source_config)
data_target['dashboards'] = search_dashboards(target_config)

# Create a new Document
doc = Document()

def analyze_additions(source, target):
    additions = {}
    for key, value in source.items():
        if key not in target:
            additions[key] = value
    return additions

def analyze_merges(source, target):
    merges = {}
    for key, value in source.items():
        if key in target:
            merges[key] = value
    return merges

# Section functions for each entity
def add_projects_section(doc, source_data, target_data):
    source_count = len(source_data)
    target_count = len(target_data)

    logging.info(f"Analyzing Projects")
    logging.info(f"Source count: {source_count}, Target count: {target_count}")
    
    additions = analyze_additions({item['key']: item.get('description', 'No description') for item in source_data}, {item['key']: item.get('description', 'No description') for item in target_data})
    conflicts = analyze_merges({item['key']: item.get('description', 'No description') for item in source_data}, {item['key']: item.get('description', 'No description') for item in target_data})
    
    doc.add_heading('Projects', level=1)
    doc.add_paragraph(f"• Number of projects in source instance: {source_count}")
    doc.add_paragraph(f"• Number of projects in target instance: {target_count}")

    doc.add_heading('Analysis of what will be added to the target instance', level=2)
    if additions:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Key'
        hdr_cells[1].text = 'Description'
        for addition_key, addition_description in additions.items():
            row_cells = table.add_row().cells
            row_cells[0].text = addition_key
            row_cells[1].text = addition_description
    else:
        doc.add_paragraph("We didn't identify anything that will be added.")

    if conflicts:
        doc.add_heading('These project keys have conflicts and will need to be renamed in the source or target instance in order to be migrated.', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Key'
        hdr_cells[1].text = 'Description'
        for conflict_key, conflict_description in conflicts.items():
            row_cells = table.add_row().cells
            row_cells[0].text = conflict_key
            row_cells[1].text = conflict_description

def add_priorities_section(doc, source_data, target_data):
    analyze_and_add_section(doc, 'Priorities', source_data, target_data, 'name')

def add_resolutions_section(doc, source_data, target_data):
    analyze_and_add_section(doc, 'Resolutions', source_data, target_data, 'name')

def add_roles_section(doc, source_data, target_data):
    analyze_and_add_section(doc, 'Project Roles', source_data, target_data, 'name')

def add_issuetypes_section(doc, source_data, target_data):
    analyze_and_add_section(doc, 'Issue Types', source_data, target_data, 'name')

def add_filters_section(doc, source_data, target_data):
    source_count = len(source_data)
    target_count = len(target_data)

    logging.info(f"Analyzing Filters")
    logging.info(f"Source count: {source_count}, Target count: {target_count}")

    doc.add_heading('Filters', level=1)
    doc.add_paragraph(f"• Number of filters in source instance: {source_count}")
    doc.add_paragraph(f"• Number of filters in target instance: {target_count}")

    doc.add_heading('Filters that have conflicts (same name in the source and target):', level=2)
    conflicts = [item['name'] for item in source_data if item['name'] in [i['name'] for i in target_data]]
    if conflicts:
        for conflict in conflicts:
            doc.add_paragraph(f"• {conflict}")
    else:
        doc.add_paragraph("No conflicts were identified regarding filters during the assessment.")

def add_dashboards_section(doc, source_data, target_data):
    source_count = len(source_data)
    target_count = len(target_data)

    logging.info(f"Analyzing Dashboards")
    logging.info(f"Source count: {source_count}, Target count: {target_count}")

    doc.add_heading('Dashboards', level=1)
    doc.add_paragraph(f"• Number of dashboards in source instance: {source_count}")
    doc.add_paragraph(f"• Number of dashboards in target instance: {target_count}")

    doc.add_heading('Dashboards that have conflicts (same name in the source and target):', level=2)
    conflicts = [item['name'] for item in source_data if item['name'] in [i['name'] for i in target_data]]
    if conflicts:
        for conflict in conflicts:
            doc.add_paragraph(f"• {conflict}")
    else:
        doc.add_paragraph("No conflicts were identified regarding Dashboards during the assessment.")

    doc.add_heading('Limitations', level=2)
    doc.add_paragraph("Since the migration of dashboards is done through REST API, our limitation is directly tied to the lack of necessary APIs for the migration. Currently, for dashboards, we have identified the following limitations:")
    doc.add_paragraph("• Inability to migrate the layout of the dashboards")
    doc.add_paragraph("• Inability to migrate the owner of the dashboard. To work around this limitation, we add the owner as an editor of the dashboard, but the owner will be the user running the script.")
    doc.add_paragraph("• Favorite dashboards of each user will be lost, requiring each user to find their dashboard in the destination instance and set it as a favorite again.")

# Special handling for custom fields to match the required format
def add_custom_fields_section(doc, source_data, target_data):
    doc.add_heading('Custom Fields', level=1)
    
    source_count = len(source_data)
    target_count = len(target_data)
    
    doc.add_paragraph(f"• Number of custom fields in source instance: {source_count}")
    doc.add_paragraph(f"• Number of custom fields in target instance: {target_count}")

    def get_field_type(field):
        return field.get('schema', {}).get('type', 'N/A')

    non_migratable_types = ["option-with-child", "project", "sd-servicelevelagreement", "multiuserpicker"]

    source_fields = {field['name']: get_field_type(field) for field in source_data}
    target_fields = {field['name']: get_field_type(field) for field in target_data}

    # Additions table
    doc.add_heading('Custom fields that will be added', level=2)
    additions = {name: source_type for name, source_type in source_fields.items() if name not in target_fields and source_type not in non_migratable_types}
    if additions:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Source Type'
        for name, source_type in additions.items():
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = source_type
    else:
        doc.add_paragraph("No custom fields will be added.")

    # Merges table
    doc.add_heading('Custom fields with the same name on both instances', level=2)
    merges = {name: source_type for name, source_type in source_fields.items() if name in target_fields and source_type != target_fields[name]}
    if merges:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Source Type'
        hdr_cells[2].text = 'Target Type'
        hdr_cells[3].text = 'Suggestion'
        for name, source_type in merges.items():
            target_type = target_fields.get(name, 'N/A')
            suggestion = 'Rename source/target or Merge or Delete'
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = source_type
            row_cells[2].text = target_type
            row_cells[3].text = suggestion
    else:
        doc.add_paragraph("No custom fields have the same name on both instances with different types.")

    # Non-migratable fields table
    doc.add_heading('Custom fields that will not be migrated', level=2)
    non_migratable = {name: source_type for name, source_type in source_fields.items() if source_type in non_migratable_types}
    if non_migratable:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Type'
        for name, source_type in non_migratable.items():
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = source_type
    else:
        doc.add_paragraph("No custom fields will not be migrated.")

# Special handling for statuses to only include ones with same name but different status categories
def add_statuses_section(doc, source_data, target_data):
    doc.add_heading('Statuses', level=1)

    source_count = len(source_data)
    target_count = len(target_data)
    
    doc.add_paragraph(f"• Number of statuses in source instance: {source_count}")
    doc.add_paragraph(f"• Number of statuses in target instance: {target_count}")

    # Additions table
    doc.add_heading('Statuses that will be added', level=2)
    additions = analyze_additions({status['name']: status['statusCategory']['name'] for status in source_data}, {status['name']: status['statusCategory']['name'] for status in target_data})
    if additions:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Category'
        for addition_name, addition_category in additions.items():
            row_cells = table.add_row().cells
            row_cells[0].text = addition_name
            row_cells[1].text = addition_category
    else:
        doc.add_paragraph("No statuses will be added.")

    # Conflicts table
    doc.add_heading('Statuses with the same name but different categories', level=2)
    source_statuses = {status['name']: status['statusCategory']['name'] for status in source_data}
    target_statuses = {status['name']: status['statusCategory']['name'] for status in target_data}
    conflicts = {name: source_category for name, source_category in source_statuses.items() if name in target_statuses and source_category != target_statuses[name]}
    if conflicts:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Source Category'
        hdr_cells[2].text = 'Target Category'
        hdr_cells[3].text = 'Suggestion'
        for name, source_category in conflicts.items():
            target_category = target_statuses[name]
            suggestion = 'Change category or Merge'
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = source_category
            row_cells[2].text = target_category
            row_cells[3].text = suggestion
    else:
        doc.add_paragraph("No statuses have the same name but different categories.")

# Function to analyze and add sections for all required entities
def analyze_and_add_section(doc, title, source_data, target_data, key_attr):
    source_count = len(source_data)
    target_count = len(target_data)

    logging.info(f"Analyzing {title}")
    logging.info(f"Source count: {source_count}, Target count: {target_count}")
    
    additions = analyze_additions({item[key_attr]: item.get('description', 'No description') for item in source_data}, {item[key_attr]: item.get('description', 'No description') for item in target_data})
    merges = analyze_merges({item[key_attr]: item.get('description', 'No description') for item in source_data}, {item[key_attr]: item.get('description', 'No description') for item in target_data})
    
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"• Number of {title.lower()} in source instance: {source_count}")
    doc.add_paragraph(f"• Number of {title.lower()} in target instance: {target_count}")

    doc.add_heading('Analysis of what will be added to the target instance', level=2)
    if additions:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Description'
        for addition_name, addition_description in additions.items():
            row_cells = table.add_row().cells
            row_cells[0].text = addition_name
            row_cells[1].text = addition_description
    else:
        doc.add_paragraph("We didn't identify anything that will be added.")

    doc.add_heading('What will be merged (because they are on both instances)', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Description'
    for merge_name, merge_description in merges.items():
        row_cells = table.add_row().cells
        row_cells[0].text = merge_name
        row_cells[1].text = merge_description

# Add sections for all required entities with error handling
sections = [
    ('Projects', 'projects', 'key'),
    ('Priorities', 'priorities', 'name'),
    ('Resolutions', 'resolutions', 'name'),
    ('Issue Types', 'issuetypes', 'name'),
    ('Filters', 'filters', 'name'),
    ('Dashboards', 'dashboards', 'name')
]

for title, key, attr in sections:
    try:
        if title == 'Projects':
            add_projects_section(doc, data_source[key], data_target[key])
        elif title == 'Filters':
            add_filters_section(doc, data_source[key], data_target[key])
        elif title == 'Dashboards':
            add_dashboards_section(doc, data_source[key], data_target[key])
        else:
            analyze_and_add_section(doc, title, data_source[key], data_target[key], attr)
    except Exception as e:
        logging.error(f"Failed to analyze {title}: {e}")

# Special handling for custom fields with error handling
try:
    add_custom_fields_section(doc, data_source['customfields'], data_target['customfields'])
except Exception as e:
    logging.error(f"Failed to analyze custom fields: {e}")

# Special handling for statuses with error handling
try:
    add_statuses_section(doc, data_source['statuses'], data_target['statuses'])
except Exception as e:
    logging.error(f"Failed to analyze statuses: {e}")

# Save the document
doc_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'jira_analysis.docx')
doc.save(doc_path)
logging.info(f"Document saved to {doc_path}")
